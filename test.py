import json
import spacy
from spacy.util import minibatch, compounding
from spacy.training import Example
import docx    #docx, pdf
import PyPDF2  #pdf-1
import fitz    #odf-2
from docx import Document  #odf-2
import subprocess
import tempfile
##from sklearn.metrics import precision_score, recall_score, f1_score

#预加载一个模型
nlp = spacy.load("en_core_web_sm")

#添加标签
ner = nlp.get_pipe("ner")
for label in ["Time", "Place", "Adjectives", "Verb", "Relationship", "Acts", "People", "Numbers", "OtherNons", "Name", "Adverbs", "Meaningless", "Item"]:
    ner.add_label(label)

#加载训练集
with open("1.jsonl", "r", encoding="utf-8") as f:
    TRAIN_DATA = []
    for line in f:
        data = json.loads(line)
        text = data["text"]
        entities = data["label"]
        # entities = data["entities"]
        annotations = {"entities": entities}
        # for entity in entities:
        #     start = entity["start"]
        #     end = entity["end"]
        #     label = entity["label"]
        # for entity in entities:
        #     start, end, label = entity["start"], entity["end"], entity["label"]


        # for entity in entities:
        #     label, start, end = entity
        #     annotations["entities"].append((start, end, label))
        TRAIN_DATA.append((text, annotations))

# 训练
n_iter = 50
other_pipes = [pipe for pipe in nlp.pipe_names if pipe != "ner"]
with nlp.disable_pipes(*other_pipes):
    optimizer = nlp.begin_training()
    for i in range(n_iter):
        losses = {}
        batches = minibatch(TRAIN_DATA, size=compounding(4.0, 32.0, 1.001))
        for batch in batches:
            texts, annotations = zip(*batch)
            examples = []
            for i in range(len(texts)):
                examples.append(Example.from_dict(nlp.make_doc(texts[i]), annotations[i]))
            nlp.update(examples, drop=0.2, sgd=optimizer, losses=losses)
        print(f"Iteration {i}: Losses - {losses}")





# docx
def read_docx(file_path):
    doc = docx.Document(file_path)
    text_test = "\n".join([para.text for para in doc.paragraphs])
    return text_test

file_path = '222.docx'
text_test = read_docx(file_path)





# # doc
# import os
# import win32com.client as win32

# # 输入doc 绝对路径
# input_file = 'E:\\homework\\Второй симестр\\АВТ.системы\\Первая задача\\word4.doc'

# # 输出
# output_file = os.path.splitext(input_file)[0] + '.docx'

# # 开始一个处理doc功能
# word = win32.gencache.EnsureDispatch('Word.Application')

# # 检查是否重名
# for doc in word.Documents:
#     if doc.Name == os.path.basename(output_file):
#         doc.Close(False)

# # 输出
# doc = word.Documents.Open(input_file)
# docx_path = os.path.join(os.path.dirname(input_file), output_file)
# doc.SaveAs(docx_path, FileFormat=16)  # 16 is the value for .docx format
# doc.Close()

# # quit Word
# word.Quit()

# def read_doc(file_path):
#     doc = docx.Document(file_path)
#     text_test = "\n".join([para.text for para in doc.paragraphs])
#     return text_test

# file_path = 'word4.docx'
# text_test = read_doc(file_path)





# # pdf-1
# def read_pdf(file_path):
#     with open(file_path, 'rb') as file:
#         pdf_reader = PyPDF2.PdfReader(file)
#         num_pages = len(pdf_reader.pages)
#         text_test = ''
#         for page in range(num_pages):
#             page_obj = pdf_reader.pages[page]
#             text_test += page_obj.extract_text()
#     return text_test

# file_path = 'example.pdf'
# text_test = read_pdf(file_path)





# # pdf-2
# pdf_file = 'word4.pdf'
# pdf_doc = fitz.open(pdf_file)

# # Initialize Word document
# docx_doc = Document()

# # Iterate over pages
# for pg in range(pdf_doc.page_count):
#     page = pdf_doc[pg]
    
#     # Add new page to Word document
#     if pg > 0:
#         docx_doc.add_page_break()
    
#     # Iterate over page text blocks
#     for blk in page.get_text("text").split('\n'):
#         # Add new paragraph to Word document
#         para = docx_doc.add_paragraph()
        
#         # Add text block to paragraph
#         para.add_run(blk)
        
# # Save Word document
# docx_file = 'example.docx'
# docx_doc.save(docx_file)

# def read_docx(file_path):
#     doc = docx.Document(file_path)
#     text_test = "\n".join([para.text for para in doc.paragraphs])
#     return text_test

# file_path = 'example.docx'
# text_test = read_docx(file_path)





# # djvu


# # open the DJVU file
# doc = fitz.open('aqeel-ur-rehman2014.djvu')

# # create a PDF document
# pdf_doc = fitz.open()

# # iterate over the pages of the DJVU file
# for page in doc:
#     # convert the page to a PDF format and add it to the PDF document
#     pdf_bytes = page.getPixmap(alpha=False).getPDFData()
#     pdf_page = fitz.Document(stream=pdf_bytes, filetype='pdf')
#     pdf_doc.insertPDF(pdf_page)

# # save the PDF document
# pdf_doc.save('document.pdf')

# # pdf-2
# pdf_file = 'word4.pdf'
# pdf_doc = fitz.open(pdf_file)

# # Initialize Word document
# docx_doc = Document()

# # Iterate over pages
# for pg in range(pdf_doc.page_count):
#     page = pdf_doc[pg]
    
#     # Add new page to Word document
#     if pg > 0:
#         docx_doc.add_page_break()
    
#     # Iterate over page text blocks
#     for blk in page.get_text("text").split('\n'):
#         # Add new paragraph to Word document
#         para = docx_doc.add_paragraph()
        
#         # Add text block to paragraph
#         para.add_run(blk)
        
# # Save Word document
# docx_file = 'document.docx'
# docx_doc.save(docx_file)

# def read_docx(file_path):
#     doc = docx.Document(file_path)
#     text_test = "\n".join([para.text for para in doc.paragraphs])
#     return text_test

# file_path = 'document.docx'
# text_test = read_docx(file_path)



doc = nlp(text_test)
print("Entities:", [(ent.text, ent.label_) for ent in doc.ents])

# # #doc = nlp("March 11, 9:28 pm, Nanjing Meteorological Station issued a yellow warning signal for cold wave: affected by strong cold air, the daily minimum temperature in most areas of Nanjing is expected to drop by 12℃ to 13℃ for 48 hours from today to 13, accompanied by northerly winds of 5 to 7 on land and 6 to 8 on water. 13 morning minimum temperature: 2℃ to 4℃. Please take precautions.")

# # predict
# predicted_entities = [(ent.text, ent.label_) for ent in doc.ents]

# # Define the actual entities (ground truth)
# actual_entities = [("entity1", "LABEL1"), ("entity2", "LABEL2"), ...]

# # Calculate precision, recall, and F1 score
# precision = precision_score(actual_entities, predicted_entities, average="weighted")
# recall = recall_score(actual_entities, predicted_entities, average="weighted")
# f1 = f1_score(actual_entities, predicted_entities, average="weighted")

# print("Precision:", precision)
# print("Recall:", recall)
# print("F1 score:", f1)
